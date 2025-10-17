import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from pathlib import Path
from datetime import datetime
import uuid
import os
import textwrap

from moviepy.editor import ImageClip, concatenate_videoclips, AudioFileClip
import pyttsx3

# ---------------------------
# Config
# ---------------------------
OUTPUT_DIR = Path("./outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

def unique_filename(prefix="out", ext="png"):
    return str(OUTPUT_DIR / f"{prefix}_{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}_{uuid.uuid4().hex[:8]}.{ext}")

# ---------------------------
# Image generator
# ---------------------------
def create_image_with_text(text="Hello World", image_size=(512,512), margin=20):
    img = Image.new("RGB", image_size, (255,255,255))
    draw = ImageDraw.Draw(img)
    font_size = 100
    try:
        font_path = "arial.ttf"
        font = ImageFont.truetype(font_path, font_size)
    except:
        font = ImageFont.load_default()
    max_width = image_size[0] - 2*margin
    max_height = image_size[1] - 2*margin

    while font_size > 10:
        try:
            font = ImageFont.truetype(font_path, font_size)
        except:
            font = ImageFont.load_default()
        lines = textwrap.wrap(text, width=40)
        line_sizes = [draw.textbbox((0,0), line, font=font) for line in lines]
        line_widths = [bbox[2]-bbox[0] for bbox in line_sizes]
        line_heights = [bbox[3]-bbox[1] for bbox in line_sizes]
        total_height = sum(line_heights)
        if max(line_widths)<=max_width and total_height<=max_height:
            break
        font_size -= 1

    avg_char_width = max(line_widths)/max(len(line) for line in lines)
    max_chars_per_line = int(max_width/avg_char_width)
    lines = textwrap.wrap(text, width=max_chars_per_line)

    line_sizes = [draw.textbbox((0,0), line, font=font) for line in lines]
    line_heights = [bbox[3]-bbox[1] for bbox in line_sizes]
    total_height = sum(line_heights)
    y_offset = (image_size[1]-total_height)/2

    for i,line in enumerate(lines):
        bbox = draw.textbbox((0,0), line, font=font)
        text_width = bbox[2]-bbox[0]
        text_height = bbox[3]-bbox[1]
        x = (image_size[0]-text_width)/2
        draw.text((x, y_offset), line, fill=(0,0,0), font=font)
        y_offset += text_height

    path = unique_filename("image","png")
    img.save(path)
    return path

# ---------------------------
# Chart generator
# ---------------------------
def create_chart(values=[5,3,7,2], labels=None):
    labels = labels or [f"Item {i}" for i in range(len(values))]
    plt.figure(figsize=(6,4))
    plt.bar(labels, values, color='skyblue')
    plt.title("Bar Chart")
    path = unique_filename("chart","png")
    plt.savefig(path)
    plt.close()
    return path

# ---------------------------
# DOCX generator
# ---------------------------
def export_docx(title="Doc", body="Hello"):
    doc = Document()
    doc.add_heading(title, level=1)
    para = doc.add_paragraph(body)
    try:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except:
        para.alignment = 0
    path = unique_filename("doc","docx")
    doc.save(path)
    return path

# ---------------------------
# PDF generator
# ---------------------------
def export_pdf(title="PDF", body="Hello"):
    path = unique_filename("doc","pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    body_style = styles['BodyText']
    body_style.alignment = 4
    story.append(Paragraph(title,title_style))
    story.append(Spacer(1,12))
    story.append(Paragraph(body.replace("\n","<br/>"), body_style))
    doc.build(story)
    return path

# ---------------------------
# TTS offline
# ---------------------------
def text_to_speech(text, audio_file="speech.mp3"):
    engine = pyttsx3.init()
    engine.setProperty('rate', 150)
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[0].id)
    engine.save_to_file(text, audio_file)
    engine.runAndWait()
    return audio_file

# ---------------------------
# Video generation with optional audio
# ---------------------------
def create_video_from_text(text="Hello World", audio_text=None, video_size=(512,512), duration_per_slide=2):
    lines = text.split("\n")
    clips = []
    for line in lines:
        img_path = create_image_with_text(line, image_size=video_size, margin=20)
        clip = ImageClip(img_path).set_duration(duration_per_slide)
        clips.append(clip)
    final_clip = concatenate_videoclips(clips, method="compose")
    if audio_text:
        audio_file = text_to_speech(audio_text, unique_filename("audio","mp3"))
        audio_clip = AudioFileClip(audio_file)
        final_clip = final_clip.set_audio(audio_clip)
    video_path = unique_filename("video","mp4")
    final_clip.write_videofile(video_path, fps=24, codec="libx264", audio=True)
    return video_path

# ---------------------------
# Streamlit app
# ---------------------------
st.set_page_config(page_title="Offline Generator", layout="wide")
st.title("Offline Generator: Image / Chart / DOCX / PDF / Video with Audio")

option = st.radio("Select what you want to generate:", ["Image","Chart","DOCX","PDF","Video"])

def download_file(file_path):
    with open(file_path,"rb") as f:
        st.download_button(label=f"Download {os.path.basename(file_path)}", data=f, file_name=os.path.basename(file_path))

# IMAGE
if option=="Image":
    st.header("Generate Image")
    img_text = st.text_area("Enter text for image:","Hello World")
    if st.button("Create Image"):
        img_file = create_image_with_text(img_text)
        st.image(img_file)
        st.success(f"Image saved at {img_file}")
        download_file(img_file)

# CHART
elif option=="Chart":
    st.header("Generate Chart")
    chart_vals = st.text_input("Enter comma-separated values:","5,3,7,2")
    chart_labels = st.text_input("Enter comma-separated labels (optional)","")
    if st.button("Create Chart"):
        try:
            values = [int(v.strip()) for v in chart_vals.split(",")]
            labels = [l.strip() for l in chart_labels.split(",")] if chart_labels else None
            chart_file = create_chart(values, labels)
            st.image(chart_file)
            st.success(f"Chart saved at {chart_file}")
            download_file(chart_file)
        except:
            st.error("Invalid input! Please enter integers separated by commas.")

# DOCX
elif option=="DOCX":
    st.header("Generate DOCX")
    doc_title = st.text_input("DOCX Title:","My Document")
    doc_body = st.text_area("DOCX Body:","Hello world!")
    if st.button("Create DOCX"):
        docx_file = export_docx(doc_title, doc_body)
        st.success(f"DOCX saved at {docx_file}")
        download_file(docx_file)

# PDF
elif option=="PDF":
    st.header("Generate PDF")
    pdf_title = st.text_input("PDF Title:","My PDF")
    pdf_body = st.text_area("PDF Body:","Hello world!")
    if st.button("Create PDF"):
        pdf_file = export_pdf(pdf_title, pdf_body)
        st.success(f"PDF saved at {pdf_file}")
        download_file(pdf_file)

# VIDEO
elif option=="Video":
    st.header("Generate Video from Text (Optional Audio)")
    video_text = st.text_area("Enter text for video slides:","Hello World\nSecond Slide")
    audio_text = st.text_area("Enter text for narration (optional):","")  # TTS
    video_duration = st.number_input("Seconds per slide:", min_value=1, max_value=20, value=2)
    if st.button("Create Video"):
        with st.spinner("Generating video..."):
            video_file = create_video_from_text(video_text, audio_text=audio_text if audio_text else None, duration_per_slide=video_duration)
        st.video(video_file)
        st.success(f"Video saved at {video_file}")
        download_file(video_file)
